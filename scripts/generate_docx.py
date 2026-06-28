#!/usr/bin/env python3
"""
Generate editable DOCX from Markdown for KORRO Research v2.

Usage:
    python generate_docx.py paper.md paper.docx
    python generate_docx.py paper.md
    python generate_docx.py --help
"""

import argparse
import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_OK = True
except ImportError:
    DOCX_OK = False


def md_to_docx(md_path: str, docx_path: str):
    if not DOCX_OK:
        print("[ERROR] python-docx not installed. Run: pip install python-docx", file=sys.stderr)
        sys.exit(1)

    text = Path(md_path).read_text(encoding='utf-8')
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)

    lines = text.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]

        # Skip empty lines
        if not line.strip():
            i += 1
            continue

        # Headers
        if line.startswith('# '):
            h = doc.add_heading(line[2:].strip(), level=1)
            i += 1
            continue
        if line.startswith('## '):
            h = doc.add_heading(line[3:].strip(), level=2)
            i += 1
            continue
        if line.startswith('### '):
            h = doc.add_heading(line[4:].strip(), level=3)
            i += 1
            continue

        # Blockquotes (prompts and notes)
        if line.startswith('> '):
            content = line[2:]
            p = doc.add_paragraph()
            p.style = doc.styles['Normal']
            run = p.add_run(content)
            run.font.italic = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(100, 100, 100)
            i += 1
            continue

        # Code blocks
        if line.startswith('```'):
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].startswith('```'):
                code_lines.append(lines[i])
                i += 1
            if code_lines:
                p = doc.add_paragraph()
                run = p.add_run('\n'.join(code_lines))
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
            i += 1  # skip closing ```
            continue

        # Tables
        if line.startswith('|') and '---' in lines[i + 1] if i + 1 < len(lines) else False:
            table_lines = []
            while i < len(lines) and lines[i].startswith('|'):
                table_lines.append(lines[i])
                i += 1

            if len(table_lines) >= 2:
                # Parse header
                header_cells = [c.strip() for c in table_lines[0].split('|')[1:-1]]
                # Skip separator line
                rows = [[c.strip() for c in r.split('|')[1:-1]] for r in table_lines[2:]]

                table = doc.add_table(rows=1 + len(rows), cols=len(header_cells))
                table.style = 'Table Grid'

                # Header
                for j, cell_text in enumerate(header_cells):
                    cell = table.rows[0].cells[j]
                    cell.text = cell_text
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.size = Pt(9)

                # Data
                for row_idx, row in enumerate(rows):
                    for col_idx, cell_text in enumerate(row):
                        if col_idx < len(header_cells):
                            cell = table.rows[row_idx + 1].cells[col_idx]
                            cell.text = cell_text
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.size = Pt(9)
            continue

        # List items
        if line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(line[2:], style='List Bullet')
            i += 1
            continue

        if re.match(r'^\d+\.\s', line):
            content = re.sub(r'^\d+\.\s', '', line)
            p = doc.add_paragraph(content, style='List Number')
            i += 1
            continue

        # Regular paragraph
        # Strip inline markdown formatting
        cleaned = re.sub(r'\*\*(.+?)\*\*', r'\1', line)
        cleaned = re.sub(r'\*(.+?)\*', r'\1', cleaned)
        cleaned = re.sub(r'`(.+?)`', r'\1', cleaned)
        cleaned = re.sub(r'\$(.+?)\$', r'\1', cleaned)

        p = doc.add_paragraph(cleaned)
        i += 1

    doc.save(docx_path)
    return docx_path


def main():
    parser = argparse.ArgumentParser(
        description='Generate editable DOCX from Markdown',
        epilog='KORRO Research v2 - PDF for reading, DOCX for editing.',
    )
    parser.add_argument('input', nargs='?', help='Markdown input file')
    parser.add_argument('output', nargs='?', help='DOCX output file')

    args = parser.parse_args()

    if not args.input:
        parser.print_help()
        print("\nExamples:")
        print("  python generate_docx.py paper.md paper.docx")
        print("  python generate_docx.py paper.md")
        return

    inp = Path(args.input)
    if not inp.exists():
        print(f"[ERROR] File not found: {args.input}", file=sys.stderr)
        sys.exit(1)

    if args.output:
        out = Path(args.output)
    else:
        out = inp.with_suffix('.docx')

    result = md_to_docx(str(inp), str(out))
    print(f"[OK] DOCX saved: {result}")


if __name__ == '__main__':
    main()
